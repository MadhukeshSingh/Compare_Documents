"""
Document Diff Service
Handles document comparison with word-level diffs and coordinate tracking.
Storage-free: uses tempfile for all intermediate files.
Supports: PDF, DOCX, DOC, TXT, MD, PY, CPP, JAVA, JS, TS, C, H, JSON, XML, CSV, HTML
"""
import hashlib
import io
import logging
import tempfile
import time
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from dataclasses import dataclass, asdict
from functools import wraps
import re

import fitz  # PyMuPDF
from PIL import Image, ImageDraw
from difflib import SequenceMatcher

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Supported extensions
# ---------------------------------------------------------------------------
PDF_EXT = {'.pdf'}
WORD_EXT = {'.docx', '.doc'}
TEXT_EXT = {'.txt', '.md'}
CODE_EXT = {'.py', '.cpp', '.java', '.js', '.ts', '.c', '.h', '.json', '.xml',
            '.csv', '.html', '.htm', '.cs', '.go', '.rb', '.rs', '.swift',
            '.kt', '.php', '.sh', '.yaml', '.yml', '.toml', '.ini', '.cfg'}
ALL_SUPPORTED = PDF_EXT | WORD_EXT | TEXT_EXT | CODE_EXT


def timing_decorator(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        t = time.time()
        try:
            result = func(*args, **kwargs)
            logger.debug(f"{func.__name__} completed in {time.time()-t:.3f}s")
            return result
        except Exception as e:
            logger.error(f"{func.__name__} failed after {time.time()-t:.3f}s: {e}")
            raise
    return wrapper


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class ImageInfo:
    image_id: str
    page_number: int
    image_index: int
    x: float
    y: float
    width: float
    height: float
    image_hash: str
    image_path: Optional[str] = None  # kept for API compat, may be None


@dataclass
class WordChange:
    change_type: str        # added | removed | replaced
    old_text: Optional[str]
    new_text: Optional[str]
    context_before: Optional[str]
    context_after: Optional[str]
    page: Optional[int]
    line: Optional[int]
    bbox: Optional[Dict[str, float]]
    page_source: Optional[str]   # doc1 | doc2 | both
    doc1_page: Optional[int]
    doc2_page: Optional[int]
    doc1_line: Optional[int]
    doc2_line: Optional[int]
    doc1_bbox: Optional[Dict[str, float]]
    doc2_bbox: Optional[Dict[str, float]]


@dataclass
class ImageChange:
    change_type: str         # added | removed
    doc1_image: Optional[ImageInfo]
    doc2_image: Optional[ImageInfo]


@dataclass
class DiffResult:
    word_changes: List[WordChange]
    image_changes: List[ImageChange]
    summary: Dict[str, int]
    doc1_pdf_bytes: Optional[bytes] = None   # PDF bytes for preview (may be None for text)
    doc2_pdf_bytes: Optional[bytes] = None
    doc1_html: Optional[str] = None          # HTML preview for DOCX/text/code
    doc2_html: Optional[str] = None
    doc1_type: str = 'text'                  # 'pdf' | 'html' | 'text'
    doc2_type: str = 'text'
    doc1_name: str = ''
    doc2_name: str = ''


# ---------------------------------------------------------------------------
# DocumentExtractor — all intermediate files go in a temp dir
# ---------------------------------------------------------------------------

class DocumentExtractor:
    """Extracts text blocks and images from various document types."""

    def __init__(self, temp_dir: Optional[Path] = None):
        """
        Args:
            temp_dir: Optional temporary directory. If None, one will be
                      created automatically and cleaned up by the caller.
        """
        if temp_dir is None:
            self._owned_tmpdir = tempfile.mkdtemp(prefix='docdiff_')
            self.temp_dir = Path(self._owned_tmpdir)
        else:
            self._owned_tmpdir = None
            self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    @timing_decorator
    def extract_pdf_content(
        self, pdf_bytes: bytes
    ) -> Tuple[List[Dict], List[ImageInfo]]:
        """Extract word-level text blocks + images from PDF bytes."""
        text_blocks: List[Dict] = []
        images: List[ImageInfo] = []

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num, page in enumerate(doc):
            page_number = page_num + 1
            words = page.get_text("words")
            for word_tuple in words:
                x0, y0, x1, y1, word_text, block_no, line_no, word_no = word_tuple
                if not word_text.strip():
                    continue
                text_blocks.append({
                    "page": page_number,
                    "text": word_text.strip(),
                    "line": line_no,
                    "word_index": word_no,
                    "bbox": {"x": x0, "y": y0, "width": x1 - x0, "height": y1 - y0},
                })

            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                img_hash = hashlib.md5(image_bytes).hexdigest()
                rects = page.get_image_rects(xref)
                if rects:
                    x, y, x1i, y1i = rects[0]
                else:
                    x, y, x1i, y1i = 0, 0, 100, 100
                images.append(ImageInfo(
                    image_id=f"p{page_number}_i{img_index}",
                    page_number=page_number,
                    image_index=img_index,
                    x=x, y=y,
                    width=x1i - x, height=y1i - y,
                    image_hash=img_hash,
                    image_path=None,
                ))
        doc.close()
        return text_blocks, images

    # ------------------------------------------------------------------
    # DOCX / DOC  (uses python-docx; no LibreOffice / Word COM needed)
    # ------------------------------------------------------------------

    @timing_decorator
    def extract_docx_content(
        self, docx_bytes: bytes
    ) -> Tuple[List[Dict], List[ImageInfo], str]:
        """
        Extract text from DOCX bytes using python-docx.
        Also produce an HTML preview using mammoth.
        Returns: (text_blocks, images, html_preview)
        """
        from docx import Document as DocxDocument

        stream = io.BytesIO(docx_bytes)
        doc = DocxDocument(stream)

        text_blocks: List[Dict] = []
        images: List[ImageInfo] = []
        line_num = 0

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            line_num += 1
            words = para.text.split()
            for w_idx, word in enumerate(words):
                if word.strip():
                    text_blocks.append({
                        "page": 1,
                        "text": word.strip(),
                        "line": line_num,
                        "word_index": w_idx,
                        "bbox": {"x": 0, "y": 0, "width": 0, "height": 0},
                    })

        # HTML preview via mammoth
        html_preview = self._docx_to_html(docx_bytes)
        return text_blocks, images, html_preview

    @staticmethod
    def _docx_to_html(docx_bytes: bytes) -> str:
        try:
            import mammoth
            result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
            return result.value
        except ImportError:
            # fallback: plain text in <pre>
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(docx_bytes))
                text = "\n".join(p.text for p in doc.paragraphs)
                return f"<pre>{_html_escape(text)}</pre>"
            except Exception:
                return "<p>Preview unavailable</p>"
        except Exception as e:
            logger.warning(f"mammoth conversion failed: {e}")
            return "<p>Preview unavailable</p>"

    # ------------------------------------------------------------------
    # TXT / MD / Code files
    # ------------------------------------------------------------------

    @timing_decorator
    def extract_text_content(
        self, file_bytes: bytes, filename: str = ""
    ) -> Tuple[List[Dict], List[ImageInfo], str]:
        """
        Extract text from plain text / code files.
        Returns (text_blocks, [], html_preview)
        """
        try:
            text = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            text = ""

        text_blocks: List[Dict] = []
        lines = text.splitlines(keepends=False)
        for line_num, line_text in enumerate(lines, start=1):
            for w_idx, word in enumerate(line_text.split()):
                if word.strip():
                    text_blocks.append({
                        "page": 1,
                        "text": word.strip(),
                        "line": line_num,
                        "word_index": w_idx,
                        "bbox": {"x": 0, "y": 0, "width": 0, "height": 0},
                    })

        html_preview = _code_to_html(text, filename)
        return text_blocks, [], html_preview


# ---------------------------------------------------------------------------
# Highlight rendering helpers
# ---------------------------------------------------------------------------

# Highlight colors (RGBA)
COLOR_REMOVED = (255, 80, 80, 160)     # red
COLOR_ADDED = (60, 200, 100, 160)      # green
COLOR_REPLACED_OLD = (255, 120, 0, 160)  # orange (old side)
COLOR_REPLACED_NEW = (60, 180, 255, 160)  # blue  (new side)


def render_pdf_page_with_highlights(
    pdf_bytes: bytes,
    page_number: int,         # 1-based
    changes: List[WordChange],
    side: str,                # 'doc1' or 'doc2'
    zoom: float = 1.5,
) -> bytes:
    """
    Render a single PDF page as PNG bytes with change highlights overlaid.
    Returns PNG bytes.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    if page_number < 1 or page_number > len(doc):
        doc.close()
        return b""

    page = doc[page_number - 1]
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    draw = ImageDraw.Draw(img, "RGBA")

    for ch in changes:
        if ch.change_type == 'removed':
            if side == 'doc1' and ch.doc1_bbox and ch.doc1_page == page_number:
                _draw_highlight(draw, ch.doc1_bbox, COLOR_REMOVED, zoom)
        elif ch.change_type == 'added':
            if side == 'doc2' and ch.doc2_bbox and ch.doc2_page == page_number:
                _draw_highlight(draw, ch.doc2_bbox, COLOR_ADDED, zoom)
        elif ch.change_type == 'replaced':
            if side == 'doc1' and ch.doc1_bbox and ch.doc1_page == page_number:
                _draw_highlight(draw, ch.doc1_bbox, COLOR_REPLACED_OLD, zoom)
            elif side == 'doc2' and ch.doc2_bbox and ch.doc2_page == page_number:
                _draw_highlight(draw, ch.doc2_bbox, COLOR_REPLACED_NEW, zoom)

    doc.close()
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _draw_highlight(
    draw: ImageDraw.ImageDraw,
    bbox: Dict[str, float],
    color: Tuple,
    zoom: float,
):
    x = bbox["x"] * zoom
    y = bbox["y"] * zoom
    w = bbox["width"] * zoom
    h = bbox["height"] * zoom
    # Pad slightly for visibility
    pad = 2
    draw.rectangle(
        [x - pad, y - pad, x + w + pad, y + h + pad],
        fill=color,
    )


def get_pdf_page_count(pdf_bytes: bytes) -> int:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    n = len(doc)
    doc.close()
    return n


def inject_html_highlights(
    base_html: str,
    changes: List[WordChange],
    side: str,
) -> str:
    """
    Inject <mark> highlights into HTML preview for DOCX/text content.
    Words are matched by text content (line-level for text, word-level for DOCX).
    """
    if not changes:
        return base_html

    # Build sets of words to highlight per type
    removed_words: set = set()
    added_words: set = set()
    replaced_old: set = set()
    replaced_new: set = set()

    for ch in changes:
        if ch.change_type == 'removed' and side == 'doc1' and ch.old_text:
            removed_words.add(ch.old_text)
        elif ch.change_type == 'added' and side == 'doc2' and ch.new_text:
            added_words.add(ch.new_text)
        elif ch.change_type == 'replaced':
            if side == 'doc1' and ch.old_text:
                replaced_old.add(ch.old_text)
            elif side == 'doc2' and ch.new_text:
                replaced_new.add(ch.new_text)

    def wrap_word(m: re.Match) -> str:
        word = m.group(0)
        clean = word.strip()
        if clean in removed_words:
            return f'<mark style="background:#ff5050;color:#fff;border-radius:3px;padding:1px 3px">{word}</mark>'
        if clean in added_words:
            return f'<mark style="background:#3cc864;color:#fff;border-radius:3px;padding:1px 3px">{word}</mark>'
        if clean in replaced_old:
            return f'<mark style="background:#ff7800;color:#fff;border-radius:3px;padding:1px 3px">{word}</mark>'
        if clean in replaced_new:
            return f'<mark style="background:#3cb4ff;color:#fff;border-radius:3px;padding:1px 3px">{word}</mark>'
        return word

    return re.sub(r'\b\S+\b', wrap_word, base_html)


# ---------------------------------------------------------------------------
# DiffEngine
# ---------------------------------------------------------------------------

class DiffEngine:
    """Compares documents at word level."""

    def __init__(self, extractor: Optional[DocumentExtractor] = None):
        self.extractor = extractor or DocumentExtractor()

    @timing_decorator
    def compare_documents(
        self,
        doc1_bytes: bytes,
        doc1_name: str,
        doc2_bytes: bytes,
        doc2_name: str,
    ) -> DiffResult:
        """Main entry point. Returns DiffResult with all change info + preview data."""

        text1, images1, preview1, type1 = self._extract(doc1_bytes, doc1_name)
        text2, images2, preview2, type2 = self._extract(doc2_bytes, doc2_name)

        word_changes = self._compare_text_word_level(text1, text2)
        image_changes = self._compare_images(images1, images2)
        summary = self._generate_summary(word_changes, image_changes)

        # For PDF previews keep original bytes; for others keep None
        ext1 = Path(doc1_name).suffix.lower()
        ext2 = Path(doc2_name).suffix.lower()

        return DiffResult(
            word_changes=word_changes,
            image_changes=image_changes,
            summary=summary,
            doc1_pdf_bytes=doc1_bytes if ext1 == '.pdf' else None,
            doc2_pdf_bytes=doc2_bytes if ext2 == '.pdf' else None,
            doc1_html=preview1 if type1 != 'pdf' else None,
            doc2_html=preview2 if type2 != 'pdf' else None,
            doc1_type=type1,
            doc2_type=type2,
            doc1_name=doc1_name,
            doc2_name=doc2_name,
        )

    def _extract(
        self, file_bytes: bytes, filename: str
    ) -> Tuple[List[Dict], List[ImageInfo], str, str]:
        """
        Returns (text_blocks, images, html_or_empty, type_str)
        type_str: 'pdf' | 'html' | 'text'
        """
        ext = Path(filename).suffix.lower()

        if ext == '.pdf':
            text, images = self.extractor.extract_pdf_content(file_bytes)
            return text, images, "", "pdf"

        if ext in WORD_EXT:
            text, images, html = self.extractor.extract_docx_content(file_bytes)
            return text, images, html, "html"

        # Plain text / code
        text, images, html = self.extractor.extract_text_content(file_bytes, filename)
        return text, images, html, "text"

    # ------------------------------------------------------------------
    # Core diff logic (unchanged from original)
    # ------------------------------------------------------------------

    def _is_word_truly_shifted(
        self,
        word: str,
        context_before: Optional[str],
        context_after: Optional[str],
        other_words: List[str],
    ) -> bool:
        if not word:
            return False
        if not context_before and not context_after:
            return False
        for idx, w in enumerate(other_words):
            if w != word:
                continue
            if context_before and (idx == 0 or other_words[idx - 1] != context_before):
                continue
            if context_after and (idx + 1 >= len(other_words) or other_words[idx + 1] != context_after):
                continue
            return True
        return False

    @timing_decorator
    def _compare_text_word_level(
        self, text1: List[Dict], text2: List[Dict]
    ) -> List[WordChange]:
        changes: List[WordChange] = []
        words1 = [b['text'] for b in text1]
        words2 = [b['text'] for b in text2]

        matcher = SequenceMatcher(None, words1, words2, autojunk=False)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                continue

            if tag == 'delete':
                for i in range(i1, i2):
                    wb = text1[i]
                    cb = words1[i - 1] if i > 0 else None
                    ca = words1[i + 1] if i + 1 < len(words1) else None
                    if self._is_word_truly_shifted(wb['text'], cb, ca, words2):
                        continue
                    changes.append(WordChange(
                        change_type='removed',
                        old_text=wb['text'], new_text=None,
                        context_before=cb, context_after=ca,
                        page=wb.get('page'), line=wb.get('line'),
                        bbox=wb.get('bbox'), page_source='doc1',
                        doc1_page=wb.get('page'), doc2_page=None,
                        doc1_line=wb.get('line'), doc2_line=None,
                        doc1_bbox=wb.get('bbox'), doc2_bbox=None,
                    ))

            elif tag == 'insert':
                for j in range(j1, j2):
                    wb = text2[j]
                    cb = words2[j - 1] if j > 0 else None
                    ca = words2[j + 1] if j + 1 < len(words2) else None
                    if self._is_word_truly_shifted(wb['text'], cb, ca, words1):
                        continue
                    changes.append(WordChange(
                        change_type='added',
                        old_text=None, new_text=wb['text'],
                        context_before=cb, context_after=ca,
                        page=wb.get('page'), line=wb.get('line'),
                        bbox=wb.get('bbox'), page_source='doc2',
                        doc1_page=None, doc2_page=wb.get('page'),
                        doc1_line=None, doc2_line=wb.get('line'),
                        doc1_bbox=None, doc2_bbox=wb.get('bbox'),
                    ))

            elif tag == 'replace':
                num_old = i2 - i1
                num_new = j2 - j1
                for k in range(max(num_old, num_new)):
                    oi = i1 + k if k < num_old else None
                    ni = j1 + k if k < num_new else None

                    if oi is not None and ni is not None:
                        ob = text1[oi]
                        nb = text2[ni]
                        ocb = words1[oi - 1] if oi > 0 else None
                        oca = words1[oi + 1] if oi + 1 < len(words1) else None
                        ncb = words2[ni - 1] if ni > 0 else None
                        nca = words2[ni + 1] if ni + 1 < len(words2) else None
                        if (self._is_word_truly_shifted(ob['text'], ocb, oca, words2) and
                                self._is_word_truly_shifted(nb['text'], ncb, nca, words1)):
                            continue
                        changes.append(WordChange(
                            change_type='replaced',
                            old_text=ob['text'], new_text=nb['text'],
                            context_before=ocb or ncb, context_after=nca or oca,
                            page=nb.get('page'), line=nb.get('line'),
                            bbox=nb.get('bbox'), page_source='doc2',
                            doc1_page=ob.get('page'), doc2_page=nb.get('page'),
                            doc1_line=ob.get('line'), doc2_line=nb.get('line'),
                            doc1_bbox=ob.get('bbox'), doc2_bbox=nb.get('bbox'),
                        ))
                    elif oi is not None:
                        ob = text1[oi]
                        cb = words1[oi - 1] if oi > 0 else None
                        ca = words1[oi + 1] if oi + 1 < len(words1) else None
                        if self._is_word_truly_shifted(ob['text'], cb, ca, words2):
                            continue
                        changes.append(WordChange(
                            change_type='removed',
                            old_text=ob['text'], new_text=None,
                            context_before=cb, context_after=ca,
                            page=ob.get('page'), line=ob.get('line'),
                            bbox=ob.get('bbox'), page_source='doc1',
                            doc1_page=ob.get('page'), doc2_page=None,
                            doc1_line=ob.get('line'), doc2_line=None,
                            doc1_bbox=ob.get('bbox'), doc2_bbox=None,
                        ))
                    else:
                        nb = text2[ni]
                        cb = words2[ni - 1] if ni > 0 else None
                        ca = words2[ni + 1] if ni + 1 < len(words2) else None
                        if self._is_word_truly_shifted(nb['text'], cb, ca, words1):
                            continue
                        changes.append(WordChange(
                            change_type='added',
                            old_text=None, new_text=nb['text'],
                            context_before=cb, context_after=ca,
                            page=nb.get('page'), line=nb.get('line'),
                            bbox=nb.get('bbox'), page_source='doc2',
                            doc1_page=None, doc2_page=nb.get('page'),
                            doc1_line=None, doc2_line=nb.get('line'),
                            doc1_bbox=None, doc2_bbox=nb.get('bbox'),
                        ))

        return changes

    @timing_decorator
    def _compare_images(
        self, images1: List[ImageInfo], images2: List[ImageInfo]
    ) -> List[ImageChange]:
        changes: List[ImageChange] = []
        map1 = {img.image_hash: img for img in images1}
        map2 = {img.image_hash: img for img in images2}
        for h, img in map1.items():
            if h not in map2:
                changes.append(ImageChange('removed', img, None))
        for h, img in map2.items():
            if h not in map1:
                changes.append(ImageChange('added', None, img))
        return changes

    def _generate_summary(
        self,
        word_changes: List[WordChange],
        image_changes: List[ImageChange],
    ) -> Dict[str, int]:
        return {
            'words_added':    sum(1 for c in word_changes if c.change_type == 'added'),
            'words_removed':  sum(1 for c in word_changes if c.change_type == 'removed'),
            'words_replaced': sum(1 for c in word_changes if c.change_type == 'replaced'),
            'images_added':   sum(1 for c in image_changes if c.change_type == 'added'),
            'images_removed': sum(1 for c in image_changes if c.change_type == 'removed'),
            'total_changes':  len(word_changes) + len(image_changes),
        }


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def convert_to_serializable(obj: Any) -> Any:
    if hasattr(obj, '__dataclass_fields__'):
        return asdict(obj)
    elif isinstance(obj, list):
        return [convert_to_serializable(i) for i in obj]
    elif isinstance(obj, dict):
        return {k: convert_to_serializable(v) for k, v in obj.items()}
    return obj


def _html_escape(text: str) -> str:
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _code_to_html(text: str, filename: str = "") -> str:
    """Convert code/text to HTML with basic line numbers."""
    lines = text.splitlines(keepends=False)
    rows = []
    for i, line in enumerate(lines, start=1):
        escaped = _html_escape(line)
        rows.append(
            f'<tr>'
            f'<td style="color:#555;text-align:right;padding-right:12px;user-select:none;'
            f'min-width:36px;font-size:12px">{i}</td>'
            f'<td style="white-space:pre;font-family:monospace;font-size:13px">{escaped}</td>'
            f'</tr>'
        )
    return (
        '<table style="border-collapse:collapse;width:100%;background:#0d1117">'
        + "".join(rows)
        + "</table>"
    )
